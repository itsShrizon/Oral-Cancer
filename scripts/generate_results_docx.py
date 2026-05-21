"""Generate RESULTS_DOCUMENTATION.docx with the specific metric set requested:
ACC, Recall, Precision, F1, Confusion Matrix, Training Batch Time, Test Inference Time,
Model Total Training Time, FLOPs, Epoch Time, Model Size (Trainable Parameters),
Memory Usage, Latency Distribution.
"""
import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
RESULTS = os.path.join(BASE, 'results')
OUT = os.path.join(RESULTS, 'RESULTS_DOCUMENTATION.docx')

MODELS = [
    ('resnet50',            'ResNet-50'),
    ('densenet121',         'DenseNet-121'),
    ('convnext_tiny',       'ConvNeXt-Tiny'),
    ('swin_t',              'Swin-T'),
    ('efficientnet_b0',     'EfficientNet-B0'),
    ('efficientnet_v2b2',   'EfficientNet-V2-B2'),
    ('efficientnet_v2b3',   'EfficientNet-V2-B3'),
    ('efficientnet_v2s',    'EfficientNet-V2-S'),
    ('inception_v3',        'Inception-V3'),
    ('custom_efficientnet_v2_baseline_recipe', 'Custom EfficientNet V2'),
]

# AttentionHub ablation runs. 'full' aliases the existing baseline_recipe folder
# so the proposed model and the ablation table both pull from the same numbers.
ABLATION_RUNS = [
    ('custom_efficientnet_v2_ablation_none',         'No attention (donor Block-4)'),
    ('custom_efficientnet_v2_ablation_bam',          'BAM only'),
    ('custom_efficientnet_v2_ablation_triplet',      'Triplet only'),
    ('custom_efficientnet_v2_ablation_kan',          'KAN only'),
    ('custom_efficientnet_v2_ablation_bam_triplet',  'BAM + Triplet'),
    ('custom_efficientnet_v2_ablation_bam_kan',      'BAM + KAN'),
    ('custom_efficientnet_v2_ablation_triplet_kan',  'Triplet + KAN'),
    ('custom_efficientnet_v2_baseline_recipe',       'Full hub (BAM + Triplet + KAN) — proposed'),
]


def load_json(folder, name):
    p = os.path.join(RESULTS, folder, name)
    if os.path.exists(p):
        with open(p) as f:
            return json.load(f)
    return None


def load_text(folder, name):
    p = os.path.join(RESULTS, folder, name)
    if os.path.exists(p):
        with open(p) as f:
            return f.read()
    return None


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    return t


def fmt(v, digits=4):
    if v is None: return 'N/A'
    if isinstance(v, float): return f"{v:.{digits}f}"
    return str(v)


def fmt_sci(v):
    if v is None: return 'N/A'
    return f"{v:.2e}"


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('Oral Cancer Classification — Results Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Multi-model comparison for dual-head oral-pathology classification "
        "(binary Benign vs Malignant + 7-class subtype: CaS, CoS, Gum, MC, OC, OLP, OT). "
        "Held-out test set: 1,646 images."
    )
    doc.add_paragraph(
        "Metrics reported per model: Accuracy, Recall, Precision, F1, Confusion Matrix, "
        "Training Batch Time, Test Inference Time, Model Total Training Time, FLOPs, "
        "Epoch Time, Model Size (Trainable Parameters), Memory Usage, Latency Distribution."
    )

    # ====================================================================
    # 1. Classification metrics — Binary
    # ====================================================================
    doc.add_heading('1. Classification Metrics — Binary (Benign vs Malignant)', 1)
    rows = []
    for folder, label in MODELS:
        cm = load_json(folder, 'classification_metrics.json')
        if not cm: continue
        b = cm['binary']
        rows.append([label, fmt(b['accuracy']), fmt(b['recall']),
                     fmt(b['precision']), fmt(b['f1_score'])])
    rows.sort(key=lambda r: -float(r[4]))
    add_table(doc, ['Model', 'Accuracy', 'Recall', 'Precision', 'F1'], rows)

    # ====================================================================
    # 2. Classification metrics — Subtype
    # ====================================================================
    doc.add_heading('2. Classification Metrics — Subtype (7-class)', 1)
    rows = []
    for folder, label in MODELS:
        cm = load_json(folder, 'classification_metrics.json')
        if not cm: continue
        s = cm['subtype']
        rows.append([label, fmt(s['accuracy']), fmt(s['recall']),
                     fmt(s['precision']), fmt(s['f1_score'])])
    rows.sort(key=lambda r: -float(r[4]))
    add_table(doc, ['Model', 'Accuracy', 'Recall', 'Precision', 'F1'], rows)

    # ====================================================================
    # 3. Model Size (Trainable Parameters) + FLOPs + Memory Usage
    # ====================================================================
    doc.add_heading('3. Model Size, FLOPs, and Memory Usage', 1)
    rows = []
    for folder, label in MODELS:
        pm = load_json(folder, 'performance_metrics.json')
        if not pm: continue
        rows.append([
            label,
            f"{pm.get('num_parameters'):,}" if pm.get('num_parameters') else 'N/A',
            f"{pm.get('model_size_mb')} MB" if pm.get('model_size_mb') else 'N/A',
            f"{pm.get('flops_gflops')} GFLOPs" if pm.get('flops_gflops') else 'N/A (thop not installed)',
            f"{pm.get('gpu_peak_mb')} MB" if pm.get('gpu_peak_mb') else 'N/A',
            f"{pm.get('cpu_rss_mb')} MB" if pm.get('cpu_rss_mb') else 'N/A',
        ])
    rows.sort(key=lambda r: int(r[1].replace(',', '')) if r[1] != 'N/A' else 0)
    add_table(doc, ['Model', 'Trainable Parameters', 'Size on Disk',
                    'FLOPs', 'GPU Peak Memory', 'CPU RSS Memory'], rows)

    # ====================================================================
    # 4. Timing metrics: batch / inference / epoch / total training
    # ====================================================================
    doc.add_heading('4. Timing Metrics', 1)
    doc.add_paragraph(
        "Training Batch Time = forward pass time averaged over 20 training batches. "
        "Test Inference Time = mean ± std of single-image forward pass (100 runs, GPU-synchronised). "
        "Epoch Time = mean wall-clock seconds per training epoch. "
        "Total Training Time = sum across all completed epochs."
    )
    rows = []
    for folder, label in MODELS:
        pm = load_json(folder, 'performance_metrics.json')
        tt = load_json(folder, 'training_time.json')
        if not pm and not tt: continue
        rows.append([
            label,
            f"{pm.get('batch_time_ms')} ms" if pm else 'N/A',
            (f"{pm.get('inference_time_ms_mean')} ± {pm.get('inference_time_ms_std')} ms"
             if pm else 'N/A'),
            f"{tt.get('avg_epoch_time_s')} s" if tt else 'N/A',
            (f"{tt.get('epochs_completed')}/{tt.get('epochs_max')}"
             if tt else 'N/A'),
            (f"{tt.get('total_training_time_min')} min ({tt.get('total_training_time_s')} s)"
             if tt else 'N/A'),
        ])
    add_table(doc, ['Model', 'Training Batch Time', 'Test Inference Time',
                    'Epoch Time', 'Epochs', 'Total Training Time'], rows)

    # ====================================================================
    # 5. Latency Distribution (single-image)
    # ====================================================================
    doc.add_heading('5. Latency Distribution (Single-Image Inference, ms)', 1)
    doc.add_paragraph(
        "Percentile statistics across up to 200 single-image forward passes on GPU. "
        "P95 and P99 show worst-case user-facing latency."
    )
    rows = []
    for folder, label in MODELS:
        pm = load_json(folder, 'performance_metrics.json')
        if not pm or 'latency_distribution' not in pm: continue
        ld = pm['latency_distribution']
        rows.append([label, ld['mean_ms'], ld['std_ms'], ld['min_ms'],
                     ld['p50_ms'], ld['p90_ms'], ld['p95_ms'], ld['p99_ms'],
                     ld['max_ms']])
    rows.sort(key=lambda r: r[1])
    add_table(doc, ['Model', 'Mean', 'Std', 'Min', 'P50', 'P90', 'P95', 'P99', 'Max'], rows)

    # ====================================================================
    # 6. Per-model section with Confusion Matrix and Latency Distribution images
    # ====================================================================
    doc.add_heading('6. Per-Model Visuals: Confusion Matrix + Latency Distribution', 1)

    for folder, label in MODELS:
        cm_path  = os.path.join(RESULTS, folder, 'confusion_matrices.png')
        lat_path = os.path.join(RESULTS, folder, 'latency_distribution.png')
        txt      = load_text(folder, 'evaluation_results.txt')
        cm_json  = load_json(folder, 'classification_metrics.json')
        pm       = load_json(folder, 'performance_metrics.json')
        tt       = load_json(folder, 'training_time.json')

        if not (cm_json or pm or tt):
            continue  # skip empty folders like vgg19

        doc.add_heading(label, 2)

        # Quick summary line
        summary_bits = []
        if cm_json:
            summary_bits.append(
                f"Binary F1 = {cm_json['binary']['f1_score']:.4f}, "
                f"Subtype F1 = {cm_json['subtype']['f1_score']:.4f}"
            )
        if pm:
            summary_bits.append(f"{pm.get('num_parameters'):,} params")
            summary_bits.append(f"size {pm.get('model_size_mb')} MB")
            summary_bits.append(f"GPU peak {pm.get('gpu_peak_mb')} MB")
        if tt:
            summary_bits.append(
                f"trained {tt.get('epochs_completed')}/{tt.get('epochs_max')} epochs "
                f"({tt.get('total_training_time_min')} min)"
            )
        if summary_bits:
            p = doc.add_paragraph()
            p.add_run('Summary: ').bold = True
            p.add_run(' | '.join(summary_bits))

        # Confusion matrix image
        if os.path.exists(cm_path):
            p = doc.add_paragraph()
            p.add_run('Confusion Matrix:').bold = True
            doc.add_picture(cm_path, width=Inches(6.5))

        # Latency distribution image
        if os.path.exists(lat_path):
            p = doc.add_paragraph()
            p.add_run('Latency Distribution:').bold = True
            doc.add_picture(lat_path, width=Inches(6.5))

        # Per-class classification report as monospace block
        if txt:
            try:
                lines = txt.splitlines()
                start = next(i for i, l in enumerate(lines) if 'Per-Class Report' in l)
                report = '\n'.join(lines[start + 1:]).strip()
                p = doc.add_paragraph()
                p.add_run('Per-Class Report (Subtype):').bold = True
                r_p = doc.add_paragraph()
                r = r_p.add_run(report)
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
            except StopIteration:
                pass

        doc.add_page_break()

    # ====================================================================
    # 7. AttentionHub Ablation (Custom EfficientNet V2)
    # ====================================================================
    avail_ablations = [(f, l) for f, l in ABLATION_RUNS
                       if load_json(f, 'classification_metrics.json') or
                          load_json(f, 'performance_metrics.json')]
    if avail_ablations:
        doc.add_heading('7. AttentionHub Ablation (Custom EfficientNet V2)', 1)
        doc.add_paragraph(
            "Each row uses the same fair training recipe as the 8 pretrained baselines "
            "(lr=1e-4, no warmup/AMP/clip, ES=15, no TTA). The only thing that changes "
            "between rows is which subset of attention branches is active at Stage 4. "
            "The 'No attention' control replaces the AttentionHub with the donor "
            "EfficientNetV2-B0 Block-4 (MBConv+SE) — i.e. the layer the hub replaces."
        )
        rows = []
        for folder, label in avail_ablations:
            cm = load_json(folder, 'classification_metrics.json')
            pm = load_json(folder, 'performance_metrics.json')
            bin_f1 = sub_f1 = 'N/A'
            if cm:
                bin_f1 = fmt(cm['binary']['f1_score'])
                sub_f1 = fmt(cm['subtype']['f1_score'])
            params = f"{pm.get('num_parameters'):,}" if pm and pm.get('num_parameters') else 'N/A'
            gflops = f"{pm.get('flops_gflops')}" if pm and pm.get('flops_gflops') else 'N/A'
            mb     = f"{pm.get('model_size_mb')}" if pm and pm.get('model_size_mb') else 'N/A'
            rows.append([label, params, gflops, mb, bin_f1, sub_f1])
        add_table(doc, ['Variant', 'Params', 'GFLOPs', 'Size (MB)',
                        'Binary F1', 'Subtype F1'], rows)

    # ====================================================================
    # 8. Notes
    # ====================================================================
    doc.add_heading('8. Notes', 1)
    notes = [
        "Custom EfficientNet V2 uses the same training recipe as the 8 pretrained baselines "
        "(lr=1e-4, no warmup/AMP/clip, ES=15, no TTA) for a fair comparison.",
        "Custom EfficientNet V2 has fewer trainable parameters than EfficientNet-B0 but more "
        "FLOPs because it drops EfficientNet-B0's late-stage conv_head/Block-6 (param-heavy, "
        "FLOP-light at 7×7 resolution) and replaces Block-4 with an AttentionHub (BAM + Triplet "
        "+ KAN) operating at higher spatial resolution — which is param-light but FLOP-heavy.",
        "VGG-19 folder is empty — no checkpoint was produced, so it is excluded from all "
        "tables.",
        "FLOPs are not reported because the 'thop' package is not installed. Install with "
        "'pip install thop' to populate the FLOPs column on rerun.",
        "Energy and CO2 figures (when shown) are rough estimates from wall-clock time; "
        "install codecarbon for hardware-measured values.",
    ]
    for n in notes:
        doc.add_paragraph(n, style='List Bullet')

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
