"""Convert docs/PAPER.md to docs/PAPER.docx with ACM-proceedings-style formatting.

Renders:
  - H1 -> Title (24pt bold, centered)
  - H2 -> Section heading (12pt bold)
  - H3 -> Subsection heading (11pt bold)
  - Bold/italic/code inline
  - Markdown tables -> docx tables with header row
  - Fenced code blocks -> monospace paragraphs
  - Horizontal rules -> page breaks
  - Bullet lists / numbered lists -> styled list paragraphs

Run:
  python scripts/generate_paper_docx.py
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC  = os.path.join(BASE, 'docs', 'PAPER.md')
OUT  = os.path.join(BASE, 'docs', 'PAPER.docx')


INLINE_RE = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')


def add_runs(paragraph, text):
    """Render markdown inline formatting into runs."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**'):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        elif token.startswith('*'):
            r = paragraph.add_run(token[1:-1]); r.italic = True
        elif token.startswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_table_from_markdown(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        add_runs(p, h.strip())
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
        set_cell_background(cell, 'D9E1F2')
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            add_runs(p, val.strip())
            for r in p.runs:
                r.font.size = Pt(9)


def parse_markdown_table(lines, i):
    """lines[i] is the header line '| col | col |'.
    lines[i+1] should be the separator. Returns (header, rows, next_i).
    """
    header_line = lines[i].strip().strip('|')
    header = [c.strip() for c in header_line.split('|')]
    rows = []
    j = i + 2
    while j < len(lines) and lines[j].strip().startswith('|'):
        row_line = lines[j].strip().strip('|')
        rows.append([c.strip() for c in row_line.split('|')])
        j += 1
    return header, rows, j


def is_table_start(lines, i):
    if i + 1 >= len(lines):
        return False
    if not lines[i].strip().startswith('|'):
        return False
    sep = lines[i + 1].strip()
    if not sep.startswith('|'):
        return False
    # separator looks like '|---|---|'
    body = sep.strip('|')
    parts = [p.strip() for p in body.split('|')]
    return all(set(p) <= set('-: ') and '-' in p for p in parts if p)


def add_heading(doc, text, level):
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(20)
        p.paragraph_format.space_after = Pt(12)
        return
    sizes = {2: 13, 3: 11, 4: 10}
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 10))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    # light grey background via paragraph border + monospace font
    r = p.add_run('\n'.join(code_lines))
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def render_markdown_to_docx(md_text: str, doc: Document):
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code blocks
        if stripped.startswith('```'):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # horizontal rule -> page break (sparingly)
        if stripped == '---':
            doc.add_paragraph()  # spacing
            i += 1
            continue

        # blank
        if not stripped:
            i += 1
            continue

        # headings
        m = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # tables
        if is_table_start(lines, i):
            header, rows, j = parse_markdown_table(lines, i)
            add_table_from_markdown(doc, header, rows)
            i = j
            continue

        # bullet list
        if stripped.startswith(('- ', '* ')):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, stripped[2:])
            i += 1
            continue

        # numbered list
        m = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_runs(p, m.group(2))
            i += 1
            continue

        # blockquote
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            r = p.add_run(stripped[2:])
            r.italic = True
            i += 1
            continue

        # normal paragraph (gather until blank line / next block element)
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt:
                break
            if re.match(r'^#{1,4}\s', nxt):
                break
            if nxt.startswith('```') or nxt.startswith('- ') or nxt.startswith('* '):
                break
            if re.match(r'^\d+\.\s', nxt):
                break
            if nxt.startswith('|') and is_table_start(lines, j):
                break
            para_lines.append(nxt)
            j += 1
        para_text = ' '.join(para_lines)
        p = doc.add_paragraph()
        add_runs(p, para_text)
        p.paragraph_format.space_after = Pt(6)
        i = j


def configure_document(doc: Document):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)


def main():
    with open(SRC, encoding='utf-8') as f:
        md = f.read()
    doc = Document()
    configure_document(doc)
    render_markdown_to_docx(md, doc)
    doc.save(OUT)
    print(f'Wrote {OUT}')
    print(f'Pages-equivalent estimate: ~{len(md) // 3500} pages of single-column rendered content.')


if __name__ == '__main__':
    main()
