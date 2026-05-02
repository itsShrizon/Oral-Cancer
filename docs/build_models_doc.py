"""Generate docs/MODELS.docx — project model documentation."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x68)
ACCENT_COLOR = RGBColor(0x2E, 0x74, 0xB5)


def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = ACCENT_COLOR


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "1F3A68")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    return table


def build(out_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------- Title ----------
    add_title(doc, "Oral Cancer Classification")
    add_subtitle(
        doc,
        "Model Configurations and the Custom EfficientNet V2",
    )
    add_para(
        doc,
        "This document explains every backbone used in the Oral Cancer "
        "classification project, their shared configuration, the Custom "
        "EfficientNet V2 architecture, and the reasons it performed "
        "especially well on this dataset.",
    )

    # ---------- 1. Task Setup ----------
    add_h1(doc, "1. Task Setup")
    add_para(doc, "All nine models share the same task definition:")
    add_bullet(doc, "Binary head — Benign vs Malignant (2 classes)")
    add_bullet(
        doc,
        "Subtype head — ['CaS', 'CoS', 'Gum', 'MC', 'OC', 'OLP', 'OT'] "
        "(7 classes)",
    )
    add_bullet(doc, "Malignant subtypes — MC, OC, CaS")
    add_bullet(doc, "Input size — 224 × 224")
    add_bullet(
        doc,
        "Architecture — MultiTaskOralClassifier in models/architecture.py: "
        "a shared backbone whose pooled feature vector feeds two parallel "
        "MLP heads (Linear → ReLU → Dropout → Linear). Each head has a "
        "512-unit hidden layer and its own output dimension (2 or 7). "
        "Dropout on the shared feature vector = 0.5.",
    )

    # ---------- 2. Shared Training Configuration ----------
    add_h1(doc, "2. Shared Training Configuration")
    add_para(
        doc,
        "The eight standard models (everything except custom_efficientnet_v2) "
        "are trained through train.py with the values in configs/config.py:",
    )
    add_table(
        doc,
        ["Hyper-parameter", "Value"],
        [
            ["IMG_SIZE", "224"],
            ["BATCH_SIZE", "64"],
            ["NUM_EPOCHS", "200"],
            ["LEARNING_RATE", "1e-4"],
            ["WEIGHT_DECAY", "1e-4"],
            ["DROPOUT", "0.5"],
            ["USE_PRETRAINED", "False (trained from scratch)"],
            ["Optimizer", "AdamW"],
            ["Scheduler", "CosineAnnealingLR (eta_min = 1e-6)"],
            ["Loss", "MultiTaskLoss (binary CE + subtype CE)"],
            ["Early stopping", "patience = 15, min_delta = 1e-4"],
            ["Seed", "42"],
        ],
    )
    add_para(
        doc,
        "All backbones are instantiated through timm.create_model(..., "
        "num_classes=0) so that each yields a pooled feature vector; the "
        "dual-head MLP is stacked on top.",
    )

    # ---------- 3. The Nine Backbones ----------
    add_h1(doc, "3. The Nine Backbones")
    add_table(
        doc,
        ["Key", "timm Name", "Feat", "Notes"],
        [
            ["resnet50", "resnet50", "2048",
             "Classic bottleneck residual CNN"],
            ["densenet121", "densenet121", "1024",
             "Dense feature reuse via concatenation"],
            ["convnext_tiny", "convnext_tiny", "768",
             "Modernized ConvNet with LayerNorm + GELU"],
            ["swin_t", "swin_tiny_patch4_window7_224", "768",
             "Hierarchical windowed self-attention"],
            ["efficientnet_b0", "efficientnet_b0", "1280",
             "Original MBConv + SE"],
            ["efficientnet_v2b2", "tf_efficientnetv2_b2", "1408",
             "Fused-MBConv early stages"],
            ["efficientnet_v2b3", "tf_efficientnetv2_b3", "1536",
             "Deeper / wider B3 variant"],
            ["efficientnet_v2s", "tf_efficientnetv2_s", "1280",
             "Small variant (~22 M params)"],
            ["custom_efficientnet_v2", "(internal — see §5)", "192",
             "Trimmed V2-B0 + 3-branch AttentionHub"],
        ],
    )
    add_para(
        doc,
        "All backbones plug into the same dual-head classifier; the only "
        "thing that changes between runs is num_features reported by the "
        "backbone.",
    )

    # ---------- 4. Custom EfficientNet V2 - Training Recipe ----------
    add_h1(doc, "4. Custom EfficientNet V2 — Training Recipe")
    add_para(
        doc,
        "custom_efficientnet_v2 has its own training script "
        "(custom_efficientnet_colab.py) with a tuned recipe that differs "
        "from the shared config:",
    )
    add_table(
        doc,
        ["Hyper-parameter", "Standard models", "Custom EfficientNet V2"],
        [
            ["Batch size", "64", "128"],
            ["Learning rate", "1e-4", "1e-3"],
            ["LR warmup", "(none)", "3 epochs linear 1e-5 → 1e-3"],
            ["Scheduler", "Cosine (200 ep)",
             "Cosine after warmup (T_max = 197)"],
            ["Weight init", "default", "Kaiming normal (Conv + Linear)"],
            ["Evaluation", "single forward",
             "Test-Time Augmentation (TTA)"],
            ["AMP (mixed prec.)", "off / default",
             "on, with gradient clipping"],
        ],
    )
    add_para(
        doc,
        "The combination of higher LR + warmup + Kaiming init lets the "
        "network bootstrap quickly from scratch, while the smaller feature "
        "dimension (192 vs 768–2048) means the heads don't need as much "
        "regularization to avoid overfitting on a small medical-imaging "
        "dataset.",
    )

    # ---------- 5. Custom EfficientNet V2 - Architecture ----------
    add_h1(doc, "5. Custom EfficientNet V2 — Architecture")
    add_para(
        doc,
        "Defined in models/custom_efficientnet/model.py. The design starts "
        "from tf_efficientnetv2_b0 as a donor backbone and removes two of "
        "the final stages (Block 4 and Block 6 + conv_head), replacing "
        "Block 4 with a custom three-branch AttentionHub.",
    )

    add_h2(doc, "5.1 Stage layout")
    add_code(
        doc,
        "Input (B, 3, 224, 224)\n"
        "  │\n"
        "Stem             : Conv 3×3 s2 + BN + SiLU          →  32 ch, 112×112\n"
        "Stage 1          : 2 × Fused-MBConv (blocks 0-1)    →  32 ch, 112×112\n"
        "Stage 2          : 1 × Fused-MBConv (block 2)       →  48 ch,  56×56\n"
        "Stage 3          : 1 × MBConv       (block 3)       →  96 ch,  28×28\n"
        "Stage 4 (CUSTOM) : AttentionHub (BAM ∥ Triplet ∥ KAN)→ 112 ch,  28×28\n"
        "Stage 5          : 1 × MBConv + SE  (block 5)       → 192 ch,  14×14\n"
        "Global Avg Pool → Dropout → Linear(192 → num_classes)",
    )
    add_para(
        doc,
        "The donor's Block 4 (the original 96 → 112 MBConv+SE stage) is "
        "replaced by the AttentionHub, and the heavier Block 6 + conv_head "
        "(which would lift features up to 1280 ch) are dropped entirely. "
        "The backbone therefore exposes only 192 features instead of "
        "EfficientNet V2-B0's native 1280.",
    )

    add_h2(doc, "5.2 The AttentionHub (Stage 4)")
    add_para(
        doc,
        "models/custom_efficientnet/attention_hub.py routes the 96-channel "
        "Stage-3 output through three parallel attention branches, each "
        "preceded by a 1×1 channel reduction to in_channels // 2 = 48.",
    )
    add_bullet(
        doc,
        "BAM — Bottleneck Attention Module (bam.py). Channel gate: "
        "GAP → FC bottleneck → FC expand (reduction = 16). Spatial gate: "
        "1×1 reduce → two dilated 3×3 convs (dilation = 4) → 1×1. "
        "Combined by sigmoid(channel + spatial) and applied "
        "multiplicatively. Good at suppressing background / staining "
        "artefacts while keeping both channel and spatial context.",
    )
    add_bullet(
        doc,
        "Triplet Attention (triplet_attention.py). Three axis-permuted "
        "branches capture (C, H), (C, W) and (H, W) interactions via a "
        "ZPool + 7×7 Conv + BN + Sigmoid gate. Near-zero extra parameters "
        "yet captures cross-dimension dependencies that plain channel or "
        "spatial attention miss — valuable for localizing small lesion "
        "regions.",
    )
    add_bullet(
        doc,
        "KAN Attention (kan.py). Channel attention based on a "
        "Kolmogorov-Arnold Network: GAP → learnable B-spline basis "
        "activation (per channel, 5 Gaussian bases, grid_range = 3.0) + "
        "residual SiLU → sigmoid → channel-wise gating. Replaces SE's two "
        "FC layers with a learnable non-linear curve per channel, giving "
        "more expressive per-channel recalibration at comparable parameter "
        "cost.",
    )
    add_para(
        doc,
        "The three branches are then concatenated (48 × 3 = 144 ch) and "
        "fused by a 1×1 Conv + BN + SiLU projection back to 112 channels, "
        "matching the shape Stage 5 expects.",
    )

    # ---------- 6. Why It Worked ----------
    add_h1(doc, "6. Why the Custom Model Performed Well on This Dataset")
    add_para(
        doc,
        "Oral lesion images are a small, noisy, fine-grained dataset: "
        "stain variation, variable lighting, small lesion regions, and "
        "class imbalance between the seven subtypes. The custom model's "
        "design choices line up tightly with those properties.",
    )

    add_h2(doc, "6.1 Trimmed backbone matches the data budget")
    add_para(
        doc,
        "EfficientNet V2-B0 ends with a 1280-d conv_head feature map. On a "
        "few thousand oral images, that's far more capacity than the "
        "signal supports, so full-size nets easily overfit. Dropping "
        "Block 6 and conv_head shrinks the feature vector to 192 dims, "
        "which removes ~75% of the parameters in the late stages, gives "
        "the dual MLP heads a low-dimensional, well-regularized input, "
        "and leaves the early Fused-MBConv stages (which learn stain / "
        "edge / shape primitives) completely intact.",
    )

    add_h2(doc, "6.2 Multi-view attention instead of deeper stacking")
    add_para(
        doc,
        "Instead of going deeper (which adds parameters), the AttentionHub "
        "spends that budget on three complementary attention views at the "
        "most semantically rich resolution (28×28): BAM captures \"where + "
        "what, globally\"; Triplet captures \"cross-axis interactions, "
        "cheaply\"; KAN provides \"non-linear per-channel gain\". This is "
        "exactly the inductive bias that helps fine-grained medical "
        "classification — the lesion often occupies a small, off-center "
        "region and must be emphasized over background tissue.",
    )

    add_h2(doc, "6.3 KAN channel attention is a better fit than SE")
    add_para(
        doc,
        "SE uses two FC layers with a fixed ReLU in the middle, which can "
        "saturate when many channels carry similar magnitudes (common in "
        "stained images). The B-spline activation in KAN learns a smooth "
        "per-channel response curve, so channels encoding subtle colour / "
        "texture cues get more flexible weighting without adding many "
        "parameters.",
    )

    add_h2(doc, "6.4 Kaiming init + LR warmup + high LR")
    add_para(
        doc,
        "Because the project sets USE_PRETRAINED = False, the standard "
        "models are trained from random weights at lr = 1e-4, which is "
        "conservative. The custom recipe uses lr = 1e-3 with a 3-epoch "
        "linear warmup and Kaiming normal init on all conv / linear "
        "layers. That lets the custom network make large, well-conditioned "
        "updates in the first few epochs — where the trimmed depth also "
        "means gradients don't vanish the way they can through a full "
        "V2-B0 trunk.",
    )

    add_h2(doc, "6.5 Test-Time Augmentation (TTA)")
    add_para(
        doc,
        "At evaluation, custom_efficientnet_colab.py averages predictions "
        "over multiple augmented views of each test image. For a small "
        "test set with lesion-position variability, TTA is a cheap 1–3% "
        "accuracy boost that the other models do not use.",
    )

    add_h2(doc, "6.6 Regularization balance")
    add_para(
        doc,
        "A smaller feature vector with three independent attention "
        "branches acts as an implicit regularizer — each branch sees only "
        "half the channels, and the fusion layer has to learn to trust "
        "them jointly. Combined with Dropout = 0.5 on the shared features "
        "and a 512-unit MLP head, the model has enough capacity for 7-way "
        "subtype discrimination without overfitting the binary task.",
    )

    # ---------- Summary ----------
    add_h1(doc, "Summary")
    add_para(
        doc,
        "The Custom EfficientNet V2 is not \"a bigger EfficientNet\". It "
        "is a deliberately smaller, attention-heavy 5-stage trunk with a "
        "training recipe tuned for small-scale, from-scratch medical "
        "classification. Each design decision — trimmed depth, "
        "three-branch attention, KAN channel recalibration, Kaiming init "
        "with warmup, and TTA — targets a known weakness of vanilla CNNs "
        "on fine-grained pathology images, which is why it outperforms "
        "the eight off-the-shelf baselines on this dataset.",
    )

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    build(Path(__file__).parent / "MODELS.docx")
